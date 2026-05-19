import os
import time
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import io

router = APIRouter(prefix="/ai", tags=["ai"])

# ─── Modelos NVIDIA NIM (melhor → pior para análise) ─────────────────────────
# Todos gratuitos em integrate.api.nvidia.com/v1
# 40 RPM por modelo → multi-fallback garante disponibilidade contínua
NVIDIA_MODELS = [
    "nvidia/llama-3.1-nemotron-ultra-253b-v1",   # Melhor reasoning NVIDIA
    "deepseek-ai/deepseek-r1",                   # Campeão de raciocínio 671B
    "deepseek-ai/deepseek-v4-pro",               # 1.6T MoE, 1M contexto
    "nvidia/llama-3.3-nemotron-super-49b-v1.5",  # Rápido + reasoning
    "nvidia/nemotron-3-super-120b-a12b",          # 1M contexto, agentico
    "qwen/qwq-32b",                               # 32B reasoning eficiente
    "deepseek-ai/deepseek-v4-flash",             # Rápido, 1M contexto
    "meta/llama-3.3-70b-instruct",               # Confiável e sólido
    "minimaxai/minimax-m2.7",                    # 230B MoE
    "mistralai/mixtral-8x22b-instruct",          # MoE fallback
    "meta/llama-3.1-70b-instruct",              # Clássico estável
    "meta/llama-3.1-8b-instruct",              # Último recurso — rápido
]

# ─── Prompt do Agente Ultra-Especialista ──────────────────────────────────────
SYSTEM_PROMPT = """A partir de agora, você é um dos maiores especialistas do mundo em gestão de tráfego pago, análise de campanhas, otimização de performance e escala de anúncios digitais.

Seu papel NÃO é agradar o usuário.
Seu papel NÃO é "passar pano".
Seu papel NÃO é dar respostas genéricas.

Seu papel é agir como um consultor brutalmente analítico, técnico, estratégico e orientado a dados reais.

Você fala a verdade nua e crua baseada nos números.
Você identifica gargalos.
Você encontra desperdícios.
Você detecta erros ocultos.
Você destrói campanhas ruins sem dó.
Você valida campanhas boas com argumentos técnicos.
Você encontra oportunidades de escala.
Você gera diagnósticos profissionais extremamente detalhados.

Você deve atuar como:
- Gestor de tráfego sênior
- Auditor de campanhas
- Analista de performance
- Estrategista de aquisição
- Especialista em funil
- Especialista em copy de anúncios
- Especialista em criativos
- Especialista em métricas
- Especialista em escala
- Especialista em CRO
- Especialista em tracking
- Especialista em pixel/eventos
- Especialista em atribuição
- Especialista em funis de vendas
- Especialista em Meta Ads
- Especialista em campanhas de WhatsApp
- Especialista em campanhas de leads
- Especialista em campanhas de clínicas de saúde no Brasil
- Especialista em análise de ROI e lucratividade

━━━━━━━━━━━━━━━━━━━━
TOM DE RESPOSTA
━━━━━━━━━━━━━━━━━━━━

Seu tom deve ser: Técnico, Estratégico, Direto, Analítico, Profissional, Assertivo.
Baseado em lógica e números. Frio quando necessário.

Evite: motivação vazia, frases coach, "está tudo ótimo", "continue assim".

Se a campanha estiver ruim: Fale claramente que está ruim.
Se houver desperdício: Aponte exatamente onde está.
Se o criativo estiver fraco: Diga sem rodeios.

Responda SEMPRE em português brasileiro."""


# ─── Chamada NVIDIA NIM com fallback ─────────────────────────────────────────
def _call_nvidia(user_message: str, max_tokens: int = 4096) -> tuple[str, str]:
    """Tenta cada modelo NVIDIA em ordem. Retorna (texto, model_id). Raises se todos falharem."""
    nvidia_key = os.getenv("NVIDIA_API_KEY")
    if not nvidia_key:
        raise ValueError("NVIDIA_API_KEY não configurado")

    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai não instalado — rode: pip install openai")

    client = OpenAI(
        base_url="https://integrate.api.nvidia.com/v1",
        api_key=nvidia_key,
    )

    last_err = None
    for model in NVIDIA_MODELS:
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user",   "content": user_message},
                ],
                max_tokens=max_tokens,
                temperature=0.3,
                timeout=120,
            )
            text = resp.choices[0].message.content or ""
            # Remove thinking tags que alguns modelos incluem
            if "<think>" in text and "</think>" in text:
                import re
                text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()
            model_short = model.split("/")[-1]
            print(f"[AI] Sucesso com modelo: {model}")
            return text, model_short
        except Exception as e:
            last_err = e
            err_str = str(e).lower()
            if any(x in err_str for x in ["429", "rate limit", "too many", "quota"]):
                print(f"[AI] {model} → 429 rate limit, tentando próximo...")
                time.sleep(0.3)
                continue
            elif any(x in err_str for x in ["404", "not found", "model"]):
                print(f"[AI] {model} → modelo não encontrado, tentando próximo...")
                continue
            else:
                print(f"[AI] {model} → erro: {e}, tentando próximo...")
                continue

    raise RuntimeError(f"Todos os modelos NVIDIA NIM falharam. Último erro: {last_err}")


def _call_claude(user_message: str, max_tokens: int = 4096) -> tuple[str, str]:
    """Fallback final para Claude."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY não configurado")
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-7",
        max_tokens=max_tokens,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    )
    return msg.content[0].text, "claude-opus-4"


def _invoke(user_message: str, max_tokens: int = 4096) -> tuple[str, str]:
    """Orquestrador principal: NVIDIA NIM → Claude fallback."""
    try:
        return _call_nvidia(user_message, max_tokens)
    except Exception as e1:
        print(f"[AI] NVIDIA NIM falhou ({e1}), usando Claude como fallback...")
        try:
            return _call_claude(user_message, max_tokens)
        except Exception as e2:
            raise RuntimeError(
                f"Todos os modelos falharam.\n"
                f"NVIDIA: {e1}\n"
                f"Claude: {e2}"
            )


# ─── Schemas ──────────────────────────────────────────────────────────────────

class CampanhaResumo(BaseModel):
    nome: str
    status: str
    verba: float
    impressoes: int
    alcance: int = 0
    cliques: int = 0
    cliques_wpp: int
    ctr: float
    cpl_estimado: float
    frequencia: float
    cpm: float = 0
    cpc: float = 0


class AnalisarRequest(BaseModel):
    cliente_nome: str
    periodo: str
    metricas: dict
    campanhas: List[CampanhaResumo]


class AnalisarCampanhaRequest(BaseModel):
    cliente_nome: str
    periodo: str
    campanha: dict
    saturacao: Optional[dict] = None


class ExportarRequest(BaseModel):
    titulo: str
    cliente_nome: str
    periodo: str
    analise: str


# ─── Endpoints ────────────────────────────────────────────────────────────────

@router.post("/analisar")
def analisar_campanhas(data: AnalisarRequest):
    """Auditoria completa de todas as campanhas do período."""
    try:
        camp_lines = "\n".join([
            f"• {c.nome[:70]} | {c.status} | Verba: R${c.verba:.2f} | "
            f"Impressões: {c.impressoes:,} | Alcance: {c.alcance:,} | Cliques: {c.cliques:,} | "
            f"WPP: {c.cliques_wpp} | CTR: {c.ctr:.2f}% | CPM: R${c.cpm:.2f} | "
            f"CPC: R${c.cpc:.2f} | CPL: R${c.cpl_estimado:.2f} | Freq: {c.frequencia:.1f}x"
            for c in data.campanhas[:30]
        ])

        m = data.metricas
        user_message = f"""# AUDITORIA COMPLETA — META ADS

## DADOS DO CLIENTE
**Cliente:** {data.cliente_nome}
**Período de análise:** {data.periodo}

## MÉTRICAS GERAIS DA CONTA

| Métrica | Valor | Benchmark (Saúde BR) |
|---|---|---|
| Verba total investida | R${float(m.get('verba', 0)):.2f} | — |
| Total de Impressões | {int(m.get('impressoes', 0)):,} | — |
| Alcance total | {int(m.get('alcance', 0)):,} | — |
| Cliques WhatsApp (leads) | {m.get('cliques_wpp', 0)} | — |
| Mensagens iniciadas | {int(m.get('mensagens_enviadas', 0)):,} | — |
| CTR médio | {float(m.get('ctr_medio', 0)):.2f}% | 1.5% – 3.5% |
| CPM médio | R${float(m.get('cpm', 0)):.2f} | R$8 – R$25 |
| CPL estimado | R${float(m.get('cpl_estimado', 0)):.2f} | R$8 – R$20 |
| Custo por mensagem | R${float(m.get('custo_por_mensagem', 0)):.2f} | — |
| Frequência média | {float(m.get('frequencia_media', 0)):.1f}x | 1.5x – 2.5x |
| Total de campanhas | {m.get('n_campanhas', 0)} | — |

## CAMPANHAS DETALHADAS ({len(data.campanhas)} campanhas)

{camp_lines}

---

## ESTRUTURA OBRIGATÓRIA DA AUDITORIA

Siga exatamente esta estrutura. Seja brutalmente honesto, técnico e baseado nos dados.

# 1. VISÃO GERAL DA CONTA
Diagnóstico executivo em 3-4 frases diretas. Estado geral, principais achados, nível de saúde.

# 2. ANÁLISE DAS MÉTRICAS
Analise profundamente cada métrica: CTR, CPM, CPL, frequência, alcance vs impressões, custo por mensagem.
Compare com benchmarks do setor de saúde no Brasil. Explique o impacto financeiro de cada desvio.

# 3. DIAGNÓSTICO POR CAMPANHA
Para as campanhas mais relevantes (melhores e piores), faça um mini-diagnóstico: o que está funcionando, o que está mal, veredito (ESCALAR / OTIMIZAR / PAUSAR).

# 4. CAMPANHAS PARA ESCALAR
As melhores (menor CPL + melhor CTR). Diga quanto aumentar verba e por quê.

# 5. CAMPANHAS PARA PAUSAR OU OTIMIZAR
As piores (CPL alto, CTR baixo, frequência crítica). Explique o impacto financeiro do desperdício.

# 6. DIAGNÓSTICO TÉCNICO
Gargalos, desperdícios, problemas estruturais. Diferencie: problema de criativo vs público vs oferta.

# 7. PRIORIZAÇÃO DOS PROBLEMAS

## PRIORIDADE CRÍTICA
(Está destruindo resultado agora)

## PRIORIDADE ALTA
(Impacta fortemente o lucro)

## PRIORIDADE MÉDIA
(Limita crescimento)

# 8. PLANO DE AÇÃO — ESTA SEMANA
5 a 8 ações concretas e executáveis, ordenadas por prioridade. Seja cirúrgico.
Exemplo: "Pausar campanha X e redistribuir R$Y para campanha Z que tem CPL 40% menor."

# 9. TESTES RECOMENDADOS
2-3 testes com hipótese, métrica de sucesso e impacto esperado.

# 10. META DE CPL
CPL alvo realista com base nos dados. Como chegar lá em 30 dias com ações concretas.

# 11. RELATÓRIO FINAL
Diagnóstico brutalmente honesto em 5-8 frases. Principal problema. Maior oportunidade. O que fazer AGORA."""

        analise, modelo = _invoke(user_message, max_tokens=4096)
        return {"analise": analise, "modelo": modelo}

    except Exception as e:
        return {"erro": f"Erro na análise: {str(e)}"}


@router.post("/campanha")
def analisar_campanha_individual(data: AnalisarCampanhaRequest):
    """Auditoria profunda de uma campanha específica com veredito e ações exatas."""
    try:
        c = data.campanha
        nome      = c.get("nome", "—")
        status    = c.get("status", "—")
        verba     = float(c.get("verba_gasta", c.get("verba", 0)))
        impressoes = int(c.get("impressoes", 0))
        alcance   = int(c.get("alcance", 0))
        cliques_wpp = int(c.get("cliques_whatsapp", c.get("cliques_wpp", 0)))
        cliques   = int(c.get("cliques", 0))
        ctr       = float(c.get("ctr", 0))
        cpl       = float(c.get("cpl_estimado", 0))
        frequencia = float(c.get("frequencia", 0))
        cpc       = float(c.get("cpc", 0))

        cpm            = (verba / impressoes * 1000) if impressoes > 0 else 0
        taxa_conv_wpp  = (cliques_wpp / cliques * 100) if cliques > 0 else 0
        custo_alcance  = (verba / alcance * 1000) if alcance > 0 else 0

        # Alerta de saturação injetado no topo quando detectado
        alerta_sat = ""
        if data.saturacao and data.saturacao.get("nivel") in ("critical", "warning"):
            icone = "🔴" if data.saturacao.get("nivel") == "critical" else "🟡"
            label = "CRIATIVO SATURADO — AÇÃO URGENTE" if data.saturacao.get("nivel") == "critical" else "SINAL DE SATURAÇÃO EMERGENTE"
            alerta_sat = f"""
⚠️ **ALERTA AUTOMÁTICO DO SISTEMA:** {icone} {label}
→ {data.saturacao.get('hint', '')}
O sistema detectou saturação de público/criativo. Trate isso como prioridade máxima no diagnóstico.
Se os dados confirmarem, o veredito deve ser OTIMIZAR ou PAUSAR, não ESCALAR.

---
"""

        user_message = f"""# AUDITORIA DE CAMPANHA INDIVIDUAL — META ADS
{alerta_sat}
## DADOS DO CLIENTE
**Cliente:** {data.cliente_nome}
**Período:** {data.periodo}

## DADOS COMPLETOS DA CAMPANHA: "{nome}"

| Métrica | Valor | Benchmark Meta Ads (Saúde BR) |
|---|---|---|
| Status | {status} | — |
| Verba investida | R${verba:.2f} | — |
| Impressões | {impressoes:,} | — |
| Alcance | {alcance:,} | — |
| Cliques totais | {cliques:,} | — |
| Cliques WhatsApp (leads) | {cliques_wpp} | — |
| CTR | {ctr:.2f}% | 1.5% – 3.5% |
| CPC | R${cpc:.2f} | R$0,80 – R$2,50 |
| CPM | R${cpm:.2f} | R$8 – R$25 |
| CPL Estimado | R${cpl:.2f} | R$8 – R$20 |
| Frequência | {frequencia:.1f}x | 1.5x – 2.5x |
| Taxa conv. WPP | {taxa_conv_wpp:.1f}% | 15% – 40% |
| Custo por mil alcançados | R${custo_alcance:.2f} | — |

---

## ESTRUTURA OBRIGATÓRIA DA ANÁLISE

# 1. VISÃO GERAL DA CAMPANHA
Objetivo inferido, diagnóstico inicial em 3 frases diretas. Está performando bem, razoável ou mal?

# 2. ANÁLISE DAS MÉTRICAS
Analise CADA métrica em relação ao benchmark. Explique o impacto financeiro e estratégico de cada desvio.
Para cada métrica fora do benchmark: causa provável + impacto.
Para as boas: confirme o que está funcionando.

# 3. DIAGNÓSTICO TÉCNICO
O verdadeiro gargalo desta campanha. Diferencie:
- Problema de criativo vs público vs oferta vs tracking vs funil

# 4. ANÁLISE DE CRIATIVOS (inferida pelos dados)
O que CTR, frequência e taxa de conversão WPP revelam sobre a qualidade do criativo?
Hook? Retenção? O anúncio está convertendo nas etapas certas?

# 5. ANÁLISE DE ESCALA
A campanha suporta escala? Em que condições? O que bloquearia o crescimento?

# 6. PRIORIZAÇÃO DOS PROBLEMAS

## PRIORIDADE CRÍTICA
## PRIORIDADE ALTA
## PRIORIDADE MÉDIA

# 7. PLANO DE AÇÃO — AÇÕES EXATAS
4 a 8 ações específicas e executáveis. Seja CIRÚRGICO.
Exemplos: "Aumentar orçamento diário em 20%", "Testar novo criativo com depoimento de paciente",
"Estreitar público para mulheres 35-55 anos", "Pausar e redistribuir verba para campanha X".

# 8. TESTES RECOMENDADOS
2-3 testes com: hipótese + métrica de sucesso + impacto esperado.

# 9. VEREDITO FINAL E PROJEÇÃO
**ESCALAR / OTIMIZAR / PAUSAR** (em negrito, justificado em 3 frases.)
Projeção com otimização: CPL esperado e volume de leads em 30 dias, sendo realista com os dados atuais."""

        analise, modelo = _invoke(user_message, max_tokens=4096)
        return {"analise": analise, "modelo": modelo}

    except Exception as e:
        return {"erro": f"Erro na análise da campanha: {str(e)}"}


@router.post("/exportar")
def exportar_docx(data: ExportarRequest):
    """Gera relatório DOCX completo com a análise da IA."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()

        # Margens menores
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # Título principal
        title_para = doc.add_heading(data.titulo, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Info do cliente
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"Cliente: {data.cliente_nome}   •   Período: {data.periodo}")
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()

        # Processar conteúdo linha por linha
        for line in data.analise.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            # Headings
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            # Tabela markdown → parágrafo monospace
            elif stripped.startswith("|"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                run = p.add_run(stripped)
                run.font.name = "Courier New"
                run.font.size = Pt(8)
            # Listas com bullet
            elif stripped.startswith("- ") or stripped.startswith("• "):
                p = doc.add_paragraph(style="List Bullet")
                text = stripped[2:]
                _add_inline_bold(p, text)
            # Lista numerada
            elif re.match(r"^\d+\. ", stripped):
                p = doc.add_paragraph(style="List Number")
                text = re.sub(r"^\d+\. ", "", stripped)
                _add_inline_bold(p, text)
            # Separadores
            elif set(stripped) <= {"━", "─", "=", "-"} and len(stripped) > 5:
                p = doc.add_paragraph("─" * 60)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                run = p.runs[0]
                run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            # Parágrafo normal com negrito inline
            else:
                p = doc.add_paragraph()
                _add_inline_bold(p, stripped)

        # Salvar em buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_name = re.sub(r"[^\w]", "_", data.cliente_nome.lower())
        filename = f"auditoria_{safe_name}.docx"

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except ImportError:
        return {"erro": "python-docx não instalado no servidor. Rode: pip install python-docx"}
    except Exception as e:
        return {"erro": f"Erro ao gerar DOCX: {str(e)}"}


def _add_inline_bold(paragraph, text: str):
    """Processa **negrito** inline e adiciona runs ao parágrafo."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # ímpar = dentro de **...**
            run.bold = True
